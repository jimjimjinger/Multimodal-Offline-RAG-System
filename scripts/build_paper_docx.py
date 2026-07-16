from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCIE = ROOT / "SCIE용"
SOURCE_MD = SCIE / "19_paper_draft.md"
OUT_DIR = SCIE / "논문"
FIG_DIR = SCIE / "산출물" / "도식"
OUT_DOCX = OUT_DIR / "원문 초안.docx"
FIG1 = FIG_DIR / "figure1_overall_architecture.png"
FIG2 = FIG_DIR / "figure2_g4_reranking.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(20, 20, 20)
MUTED = RGBColor(95, 95, 95)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GRID = "B8C2CC"


def font_path(name: str) -> str:
    candidates = [
        Path("C:/Windows/Fonts") / name,
        Path("C:/Windows/Fonts/malgun.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return str(path)
    return ""


FONT_REG = font_path("malgun.ttf")
FONT_BOLD = font_path("malgunbd.ttf")


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    path = FONT_BOLD if bold and FONT_BOLD else FONT_REG
    if path:
        return ImageFont.truetype(path, size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        test = word if not current else f"{current} {word}"
        bbox = draw.textbbox((0, 0), test, font=font)
        if bbox[2] - bbox[0] <= max_width:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    details: list[str],
    fill: tuple[int, int, int] = (255, 255, 255),
    outline: tuple[int, int, int] = (67, 101, 132),
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=3)
    title_font = pil_font(26, bold=True)
    detail_font = pil_font(21)
    title_lines = wrap_text(draw, title, title_font, x2 - x1 - 44)
    detail_lines: list[str] = []
    for item in details:
        detail_lines.extend(wrap_text(draw, item, detail_font, x2 - x1 - 44))
    all_lines = title_lines + detail_lines
    line_heights = []
    for idx, line in enumerate(all_lines):
        font = title_font if idx < len(title_lines) else detail_font
        bbox = draw.textbbox((0, 0), line, font=font)
        line_heights.append(bbox[3] - bbox[1])
    total = sum(line_heights) + 9 * (len(all_lines) - 1) + (10 if detail_lines else 0)
    y = y1 + ((y2 - y1) - total) // 2
    for idx, line in enumerate(all_lines):
        font = title_font if idx < len(title_lines) else detail_font
        color = (20, 48, 78) if idx < len(title_lines) else (45, 45, 45)
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) // 2, y), line, font=font, fill=color)
        y += line_heights[idx] + (18 if idx == len(title_lines) - 1 and detail_lines else 9)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int]) -> None:
    draw.line([start, end], fill=(65, 84, 102), width=5)
    ex, ey = end
    sx, sy = start
    if abs(ex - sx) >= abs(ey - sy):
        direction = 1 if ex > sx else -1
        points = [(ex, ey), (ex - 18 * direction, ey - 11), (ex - 18 * direction, ey + 11)]
    else:
        direction = 1 if ey > sy else -1
        points = [(ex, ey), (ex - 11, ey - 18 * direction), (ex + 11, ey - 18 * direction)]
    draw.polygon(points, fill=(65, 84, 102))


def draw_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    details: list[str],
    fill: tuple[int, int, int] = (247, 250, 252),
    outline: tuple[int, int, int] = (67, 101, 132),
    title_size: int = 22,
    detail_size: int = 17,
    radius: int = 16,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=3)
    title_font = pil_font(title_size, bold=True)
    detail_font = pil_font(detail_size)
    title_lines = wrap_text(draw, title, title_font, x2 - x1 - 32)
    detail_lines: list[str] = []
    for item in details:
        detail_lines.extend(wrap_text(draw, item, detail_font, x2 - x1 - 34))
    lines = title_lines + detail_lines
    heights: list[int] = []
    for idx, line in enumerate(lines):
        font = title_font if idx < len(title_lines) else detail_font
        bbox = draw.textbbox((0, 0), line, font=font)
        heights.append(bbox[3] - bbox[1])
    total = sum(heights) + 7 * max(0, len(lines) - 1) + (8 if detail_lines else 0)
    y = y1 + ((y2 - y1) - total) // 2
    for idx, line in enumerate(lines):
        font = title_font if idx < len(title_lines) else detail_font
        color = (20, 48, 78) if idx < len(title_lines) else (45, 45, 45)
        bbox = draw.textbbox((0, 0), line, font=font)
        draw.text((x1 + ((x2 - x1) - (bbox[2] - bbox[0])) // 2, y), line, font=font, fill=color)
        y += heights[idx] + (14 if idx == len(title_lines) - 1 and detail_lines else 7)


def draw_label(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], text: str, fill: tuple[int, int, int]) -> None:
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=(195, 208, 222), width=2)
    font = pil_font(24, bold=True)
    bbox = draw.textbbox((0, 0), text, font=font)
    draw.text((box[0] + 24, box[1] + ((box[3] - box[1]) - (bbox[3] - bbox[1])) // 2), text, font=font, fill=(20, 48, 78))


def create_figure_1() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (2200, 1450), "white")
    draw = ImageDraw.Draw(img)
    title_font = pil_font(38, bold=True)
    subtitle_font = pil_font(23)

    draw.text((80, 55), "Fig. 1. Overall Architecture of the Context-Aware Multimodal RAG Framework", font=title_font, fill=(20, 48, 78))
    draw.text(
        (80, 105),
        "The system separates one-time offline preprocessing from runtime retrieval, local generation, and evaluation.",
        font=subtitle_font,
        fill=(80, 80, 80),
    )

    draw_label(draw, (80, 180, 1120, 240), "Offline preprocessing layer (one-time processing before app runtime)", (236, 244, 251))
    draw_label(draw, (80, 805, 760, 865), "Runtime retrieval and response layer", (236, 244, 251))

    manual = (80, 375, 350, 555)
    text_extract = (460, 285, 760, 445)
    image_extract = (460, 500, 760, 660)
    bge = (870, 285, 1170, 445)
    siglip = (870, 500, 1170, 660)
    chroma = (1280, 285, 1580, 445)
    mapping = (1280, 500, 1580, 660)
    dataset = (1710, 375, 2070, 555)

    draw_box(draw, manual, "Robot Training Manual PDF", [], title_size=23, detail_size=17)
    draw_box(draw, text_extract, "Text Extraction and Chunking", [], fill=(236, 244, 251), title_size=22)
    draw_box(draw, image_extract, "Image Extraction and Refinement", [], fill=(236, 244, 251), title_size=22)
    draw_box(draw, bge, "BGE-M3 Text Embeddings", [], title_size=22)
    draw_box(draw, siglip, "SigLIP Features", ["Image/Text-Image Mapping"], title_size=22, detail_size=17)
    draw_box(draw, chroma, "ChromaDB Text Collection", [], fill=(242, 246, 249), title_size=22)
    draw_box(draw, mapping, "Image Metadata and Mapping Table", [], fill=(242, 246, 249), title_size=22)
    draw_box(draw, dataset, "Processed Dataset", ["text chunks + images + labels"], fill=(236, 244, 251), title_size=22, detail_size=17)

    draw_arrow(draw, (350, 455), (460, 365))
    draw_arrow(draw, (350, 475), (460, 580))
    draw_arrow(draw, (760, 365), (870, 365))
    draw_arrow(draw, (760, 580), (870, 580))
    draw_arrow(draw, (1170, 365), (1280, 365))
    draw_arrow(draw, (1170, 580), (1280, 580))
    draw_arrow(draw, (1580, 365), (1710, 455))
    draw_arrow(draw, (1580, 580), (1710, 475))

    note_font = pil_font(19)
    draw.rounded_rectangle((80, 710, 2070, 770), radius=14, fill=(250, 250, 250), outline=(210, 218, 226), width=2)
    draw.text((105, 730), "runtime uses preprocessed DB", font=note_font, fill=(60, 60, 60))
    query = (80, 990, 310, 1145)
    groups = (410, 925, 760, 1210)
    evidence = (870, 990, 1190, 1145)
    llm = (1300, 990, 1620, 1145)
    response = (1730, 990, 2070, 1145)

    draw_box(draw, query, "User Query", [], title_size=24)
    draw_box(
        draw,
        groups,
        "Evaluation Group",
        ["G1 Keyword Search", "G2 Text-only RAG", "G3 Multimodal RAG", "G4 Context-aware Re-ranking"],
        fill=(236, 244, 251),
        title_size=23,
        detail_size=17,
    )
    draw_box(draw, evidence, "Top-k Evidence", ["text evidence + image evidence"], title_size=23, detail_size=17)
    draw_box(draw, llm, "Quantized Local LLM", ["Qwen / Gemma / Llama", "4-bit runtime target"], title_size=23, detail_size=17)
    draw_box(draw, response, "Training Guidance Response", [], fill=(242, 246, 249), title_size=23)

    draw_arrow(draw, (310, 1068), (410, 1068))
    draw_arrow(draw, (760, 1068), (870, 1068))
    draw_arrow(draw, (1190, 1068), (1300, 1068))
    draw_arrow(draw, (1620, 1068), (1730, 1068))

    draw.rounded_rectangle((80, 1280, 2070, 1365), radius=14, fill=(255, 250, 235), outline=(190, 150, 65), width=2)
    draw.text(
        (105, 1302),
        "Design feature: offline operation with local vector DB and local quantized LLM; 8GB-RAM class device is a deployment target, not a fully validated benchmark yet.",
        font=pil_font(20),
        fill=(70, 60, 45),
    )
    img.save(FIG1, quality=95)


def create_figure_2() -> None:
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (2200, 1250), "white")
    draw = ImageDraw.Draw(img)
    title_font = pil_font(38, bold=True)
    subtitle_font = pil_font(23)

    draw.text((80, 55), "Fig. 2. G4 Context-Aware Re-Ranking Process", font=title_font, fill=(20, 48, 78))
    draw.text(
        (80, 105),
        "G4 is a re-ranking layer on top of G3 candidates. It adjusts candidate scores using estimated practice-stage context.",
        font=subtitle_font,
        fill=(80, 80, 80),
    )

    y = 350
    w = 265
    h = 190
    gap = 25
    xs = [80 + i * (w + gap) for i in range(7)]
    boxes = [
        ("User Query", []),
        ("G3 Candidates", ["text candidates", "image candidates", "base scores"]),
        ("Stage Estimation", ["compare query", "with stage profiles"]),
        ("Stage Context Map", ["page range", "section heading", "keywords"]),
        ("Context Score", ["page match", "section match", "keyword match"]),
        ("Score Update", ["final score = base score", "+ lambda * context score"]),
        ("Re-ranked Top-k Evidence", ["text evidence", "image evidence"]),
    ]
    for idx, x in enumerate(xs):
        draw_box(
            draw,
            (x, y, x + w, y + h),
            boxes[idx][0],
            boxes[idx][1],
            fill=(236, 244, 251) if idx in (2, 3, 4, 5) else (247, 250, 252),
            title_size=22,
            detail_size=16,
        )
        if idx < len(xs) - 1:
            draw_arrow(draw, (x + w, y + h // 2), (xs[idx + 1], y + h // 2))

    draw.rounded_rectangle((210, 760, 970, 930), radius=18, fill=(255, 250, 235), outline=(190, 150, 65), width=2)
    draw.text((245, 790), "Low-confidence fallback", font=pil_font(25, bold=True), fill=(95, 70, 10))
    draw.text((245, 835), "If stage estimation is low-confidence or ambiguous,", font=pil_font(20), fill=(60, 60, 60))
    draw.text((245, 872), "G4 keeps the original G3 ranking.", font=pil_font(20), fill=(60, 60, 60))

    draw.rounded_rectangle((1120, 760, 1990, 930), radius=18, fill=(244, 246, 249), outline=(185, 195, 205), width=2)
    draw.text((1155, 790), "No label leakage", font=pil_font(25, bold=True), fill=(20, 48, 78))
    draw.text((1155, 835), "The context map does not include question IDs,", font=pil_font(20), fill=(60, 60, 60))
    draw.text((1155, 872), "ground-truth image filenames, or answer chunk IDs.", font=pil_font(20), fill=(60, 60, 60))

    img.save(FIG2, quality=95)


def set_run_font(run, size: float | None = None, color: RGBColor | None = None, bold: bool | None = None, italic: bool | None = None, name: str = "Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_style(paragraph, before: float = 0, after: float = 6, line_spacing: float = 1.10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), GRID)


def set_table_width(table, width_dxa: int = 9360) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=MUTED)


def next_numbering_id(numbering, tag_name: str, attr_name: str) -> int:
    ids = []
    for child in numbering:
        if child.tag == qn(tag_name):
            value = child.get(qn(attr_name))
            if value is not None:
                try:
                    ids.append(int(value))
                except ValueError:
                    pass
    return (max(ids) + 1) if ids else 1


def create_numbering(doc: Document, kind: str = "decimal") -> int:
    numbering = doc.part.numbering_part.element
    abstract_id = next_numbering_id(numbering, "w:abstractNum", "w:abstractNumId")
    num_id = next_numbering_id(numbering, "w:num", "w:numId")

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))

    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")

    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)

    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)

    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)

    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    lvl.append(lvl_jc)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    p_pr.append(ind)
    lvl.append(p_pr)

    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), "Symbol")
        r_fonts.set(qn("w:hAnsi"), "Symbol")
        r_pr.append(r_fonts)
        lvl.append(r_pr)

    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id_el = num_pr.find(qn("w:numId"))
    if num_id_el is None:
        num_id_el = OxmlElement("w:numId")
        num_pr.append(num_id_el)
    num_id_el.set(qn("w:val"), str(num_id))


def add_list_paragraph(doc: Document, text: str, num_id: int) -> None:
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    set_paragraph_style(p, after=4, line_spacing=1.167)
    add_rich_text(p, text)


def add_manual_list_paragraph(doc: Document, marker: str, text: str) -> None:
    # Word's built-in list continuation can be unstable across COM/PDF export on
    # this Windows environment. Preserve source list numbers explicitly for this
    # draft so the rendered manuscript matches the Markdown source.
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_style(p, after=4, line_spacing=1.167)
    marker_run = p.add_run(f"{marker} ")
    set_run_font(marker_run, size=11)
    add_rich_text(p, text)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = MUTED
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "Context-Aware Multimodal RAG for Collaborative Robot Training"
    set_paragraph_style(header, after=0)
    for run in header.runs:
        set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_rich_text(paragraph, text: str) -> None:
    # Handles a small subset of Markdown inline syntax used in the draft.
    token_pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in token_pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size=11)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=11, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=10.5, name="Consolas", color=RGBColor(70, 70, 70))
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=11)


def add_para(doc: Document, text: str, style: str | None = None, align=None) -> None:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_style(p)
    add_rich_text(p, text)


def add_table_from_rows(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    if len(rows[0]) == 11 and "Text R@1" in rows[0][1]:
        text_rows = [[row[i] for i in [0, 1, 2, 3, 4]] for row in rows]
        image_rows = [[row[i] for i in [0, 5, 6, 7, 8, 9, 10]] for row in rows if row[0] == rows[0][0] or row[0].startswith("G3") or row[0].startswith("G4")]
        add_para(doc, "Text retrieval performance", style=None)
        add_table_from_rows(doc, text_rows)
        doc.add_page_break()
        add_para(doc, "Image and multimodal retrieval performance", style=None)
        add_table_from_rows(doc, image_rows)
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    set_table_width(table)
    set_table_borders(table)

    cols = len(rows[0])
    total_width = 6.5
    if cols == 2:
        widths = [2.1, 4.4]
    elif cols == 3:
        widths = [1.7, 2.4, 2.4]
    elif cols == 4:
        widths = [1.3, 1.7, 1.7, 1.8]
    elif cols == 5:
        widths = [2.1, 1.1, 1.1, 1.1, 1.1]
    elif cols == 6:
        widths = [1.0, 0.75, 0.75, 0.75, 0.75, 2.5]
    elif cols == 7:
        widths = [2.0, 0.75, 0.75, 0.75, 0.85, 0.75, 0.75]
    else:
        widths = [total_width / cols] * cols
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            cell.width = Inches(widths[j])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if i == 0:
                shade_cell(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j > 0 and len(value) < 18 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_style(p, after=0, line_spacing=1.10)
            run = p.add_run(value.strip())
            set_run_font(run, size=9.5 if cols >= 6 else 10, bold=(i == 0))

    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if re.fullmatch(r"\|[\s:\-|]+\|", line):
            i += 1
            continue
        row = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(row)
        i += 1
    return rows, i


def add_figure(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_style(p, before=4, after=2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.25))
    cap = doc.add_paragraph(style="Caption")
    cap.add_run(caption)


def add_algorithm(doc: Document) -> None:
    title = doc.add_paragraph()
    set_paragraph_style(title, before=8, after=4)
    r = title.add_run("Algorithm 1. Context-aware multimodal re-ranking in G4")
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)

    lines = [
        "Input: user query q, G3 text candidates T, G3 image candidates I, stage context map M",
        "Output: re-ranked text and image evidence E",
        "1: Encode q using BGE-M3.",
        "2: Compare q with stage context profiles and estimate the most relevant stage s.",
        "3: If the confidence score or top-1/top-2 margin is below the threshold, return the G3 ranking.",
        "4: Retrieve the page ranges, section headings, and keywords associated with s from M.",
        "5: For each candidate c in T and I, compute page range, section, and keyword context scores.",
        "6: Combine the G3 retrieval score and the stage context score.",
        "7: Sort candidates by the final score and return Top-k text evidence and image evidence.",
        "Constraint: question IDs, correct image filenames, and correct chunk IDs are not used in M.",
    ]
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    set_table_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, "F8FAFC")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for idx, line in enumerate(lines):
        if idx:
            p.add_run("\n")
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9.2, color=RGBColor(45, 45, 45))
    doc.add_paragraph()


def clean_markdown_line(line: str) -> str:
    return line.rstrip().replace("  ", " ")


def build_docx() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    create_figure_1()
    create_figure_2()

    source = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("A Context-Aware Multimodal Retrieval-Augmented Generation Framework\nfor Collaborative Robot Training")
    set_run_font(run, size=18, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("국문 작업 초안 | IEEE Access 투고용 구조")
    set_run_font(run, size=10.5, color=MUTED)

    # Keep only manuscript sections and references. Internal notes are excluded.
    wanted: list[str] = []
    capture = False
    capture_refs = False
    for line in source:
        if line.startswith("## 초록"):
            capture = True
        if line.startswith("## 9. 현재 초안"):
            capture = False
        if line.startswith("## 10. 참고문헌"):
            capture_refs = True
            wanted.append("## 참고문헌")
            continue
        if line.startswith("## 11. 논문에 사용할 핵심 파일"):
            capture_refs = False
        if capture or capture_refs:
            if line.startswith(">"):
                continue
            wanted.append(line)

    i = 0
    current_heading = ""
    fig1_inserted = False
    fig2_inserted = False
    algorithm_inserted = False

    while i < len(wanted):
        raw = wanted[i]
        line = clean_markdown_line(raw)
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            if current_heading.startswith("3.1") and not fig1_inserted:
                add_figure(
                    doc,
                    FIG1,
                    "Fig. 1. Overall architecture of the proposed context-aware multimodal RAG framework.",
                )
                fig1_inserted = True
            if current_heading.startswith("3.4") and not fig2_inserted:
                add_figure(
                    doc,
                    FIG2,
                    "Fig. 2. Context-aware re-ranking strategy used in G4.",
                )
                fig2_inserted = True
            if current_heading.startswith("3.4") and not algorithm_inserted:
                add_algorithm(doc)
                algorithm_inserted = True
            current_heading = stripped[4:].strip()
            doc.add_heading(current_heading, level=2)
            i += 1
            continue

        if stripped.startswith("## "):
            if current_heading.startswith("3.1") and not fig1_inserted:
                add_figure(
                    doc,
                    FIG1,
                    "Fig. 1. Overall architecture of the proposed context-aware multimodal RAG framework.",
                )
                fig1_inserted = True
            if current_heading.startswith("3.4") and not fig2_inserted:
                add_figure(
                    doc,
                    FIG2,
                    "Fig. 2. Context-aware re-ranking strategy used in G4.",
                )
                fig2_inserted = True
            if current_heading.startswith("3.4") and not algorithm_inserted:
                add_algorithm(doc)
                algorithm_inserted = True
            current_heading = stripped[3:].strip()
            doc.add_heading(current_heading, level=1)
            i += 1
            continue

        if stripped.startswith("|"):
            rows, next_i = parse_table(wanted, i)
            add_table_from_rows(doc, rows)
            i = next_i
            continue

        if re.match(r"^\d+\.\s+", stripped):
            while i < len(wanted):
                candidate = clean_markdown_line(wanted[i]).strip()
                m = re.match(r"^(\d+\.)\s+(.+)", candidate)
                if not m:
                    break
                add_manual_list_paragraph(doc, m.group(1), m.group(2))
                i += 1
            continue

        if stripped.startswith("- "):
            while i < len(wanted):
                candidate = clean_markdown_line(wanted[i]).strip()
                if not candidate.startswith("- "):
                    break
                add_manual_list_paragraph(doc, "•", candidate[2:])
                i += 1
            continue

        add_para(doc, stripped)
        i += 1

    if not fig1_inserted:
        add_figure(doc, FIG1, "Fig. 1. Overall architecture of the proposed context-aware multimodal RAG framework.")
    if not fig2_inserted:
        add_figure(doc, FIG2, "Fig. 2. Context-aware re-ranking strategy used in G4.")
    if not algorithm_inserted:
        add_algorithm(doc)

    doc.core_properties.title = "A Context-Aware Multimodal Retrieval-Augmented Generation Framework for Collaborative Robot Training"
    doc.core_properties.author = "Jimin"
    doc.core_properties.subject = "SCIE/IEEE Access manuscript draft"
    doc.core_properties.keywords = "multimodal RAG, collaborative robot training, context-aware retrieval"
    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_docx()
    print(OUT_DOCX)
