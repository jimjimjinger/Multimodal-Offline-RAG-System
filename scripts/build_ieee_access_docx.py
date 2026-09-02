from __future__ import annotations

import csv
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCIE = ROOT / "SCIE용"
SOURCE = SCIE / "논문" / "IEEE Access 영문 통합 원고.md"
TEMPLATE = SCIE / "논문" / "templates" / "Access-Template-2024.docx"
OUTPUT = SCIE / "논문" / "IEEE Access 영문 통합 원고.docx"
FIG1 = SCIE / "산출물" / "도식" / "figure1_overall_architecture.png"
FIG2 = SCIE / "산출물" / "도식" / "figure2_g4_reranking.png"
FINAL_METRICS = SCIE / "data" / "15_g1_g2_g3_g4_summary.csv"
MAPPING_ABLATION = SCIE / "data" / "32_bbox_siglip_ablation_summary.csv"
BOOTSTRAP_RESULTS = SCIE / "data" / "31_g3_g4_paired_bootstrap_ci.csv"

IEEE_BLUE = RGBColor(0x00, 0x62, 0x9B)
GRAY = RGBColor(0x58, 0x59, 0x5B)

AUTHOR_NAME = "Jimin Lee and Hyun Jung Kim"
JIMIN_AFFILIATION = (
    "Department of Artificial Intelligence, Graduate School of Information and Telecommunications, "
    "Konkuk University, Seoul 05029, Republic of Korea"
)
HYUN_JUNG_KIM_AFFILIATION = (
    "Sang-Huh College and Department of Artificial Intelligence Convergence, Graduate School of "
    "Engineering, Konkuk University, Seoul, Republic of Korea"
)

PAGINATION_MARK_TAGS = (
    "w:keepNext",
    "w:keepLines",
    "w:pageBreakBefore",
    "w:suppressLineNumbers",
    "w:widowControl",
)


def remove_paragraph_pagination_marks(doc: Document) -> None:
    """Remove paragraph pagination settings that Word displays as black squares."""
    for root in (doc.element, doc.styles.element):
        for paragraph_properties in root.xpath(".//w:pPr"):
            for tag in PAGINATION_MARK_TAGS:
                for element in paragraph_properties.findall(qn(tag)):
                    paragraph_properties.remove(element)


def ensure_dynamic_page_number(doc: Document) -> None:
    """Replace the template's static footer page number with a PAGE field."""
    footer = doc.sections[0].footer
    for paragraph in footer.paragraphs:
        if not paragraph.text.rstrip().endswith("\t1"):
            continue

        first_run = paragraph.runs[0]
        first_run.text = paragraph.text.rsplit("\t", 1)[0]
        for run in list(paragraph.runs[1:]):
            paragraph._p.remove(run._r)

        tab_run = paragraph.add_run("\t")
        tab_run.font.name = "Helvetica"
        tab_run.font.size = Pt(6)

        field_run = paragraph.add_run()
        field_run.font.name = "Helvetica"
        field_run.font.size = Pt(6)
        begin = OxmlElement("w:fldChar")
        begin.set(qn("w:fldCharType"), "begin")
        instruction = OxmlElement("w:instrText")
        instruction.set(qn("xml:space"), "preserve")
        instruction.text = " PAGE \\* MERGEFORMAT "
        separate = OxmlElement("w:fldChar")
        separate.set(qn("w:fldCharType"), "separate")
        result = OxmlElement("w:t")
        result.text = "1"
        end = OxmlElement("w:fldChar")
        end.set(qn("w:fldCharType"), "end")
        field_run._r.extend([begin, instruction, separate, result, end])
        break
AUTHOR_EMAIL = "jiminlee0508@naver.com"
AUTHOR_ORCID = "0009-0009-3159-6517"
CORRESPONDING_AUTHOR_NAME = "Hyun Jung Kim"
CORRESPONDING_AUTHOR_EMAIL = "nygirl@konkuk.ac.kr"
CORRESPONDING_AUTHOR_ORCID = "0000-0003-3845-0560"
AUTHOR_BIOGRAPHY = (
    "Jimin Lee received the B.S. degree in environmental engineering from Inha "
    "University, Incheon, Republic of Korea. Jimin Lee is currently pursuing the "
    "master's degree in artificial intelligence at the Graduate School of "
    "Information and Telecommunications, Konkuk University, Seoul, Republic of "
    "Korea. Research interests include multimodal retrieval-augmented generation, "
    "local large language models, and collaborative robot training systems."
)
CORRESPONDING_AUTHOR_BIOGRAPHY = (
    "Hyun Jung Kim is an Assistant Professor with Sang-Huh College and the "
    "Department of Artificial Intelligence Convergence, Graduate School of "
    "Engineering, Konkuk University, Seoul, Republic of Korea. Her research "
    "interests include artificial intelligence, machine learning, and efficient "
    "neural-network models."
)


MAIN_HEADINGS = {
    "1. Introduction": "INTRODUCTION",
    "2. Related Work": "RELATED WORK",
    "3. Proposed Framework": "PROPOSED FRAMEWORK",
    "4. Experimental Design": "EXPERIMENTAL DESIGN",
    "5. Experimental Results": "EXPERIMENTAL RESULTS",
    "6. Secondary Response-Quality Analysis": "SECONDARY RESPONSE-QUALITY ANALYSIS",
    "7. Discussion": "DISCUSSION",
    "8. Conclusion": "CONCLUSION",
}

TABLE_TITLES = [
    "COMPARISON GROUPS",
    "RELEVANCE CRITERIA FOR RETRIEVAL EVALUATION",
    "COMPARISON OF BBOX- AND SIGLIP-BASED TEXT-IMAGE MAPPING",
    "EXPERIMENTAL ENVIRONMENT",
    "LOCAL LANGUAGE MODELS",
    "RETRIEVAL PERFORMANCE OF G1-G4",
    "RESPONSE-QUALITY RESULTS",
]


def set_columns(section, count: int, space: int = 400) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sect_pr.append(cols)
    cols.set(qn("w:num"), str(count))
    cols.set(qn("w:space"), str(space))
    if count == 1:
        cols.attrib.pop(qn("w:equalWidth"), None)


def set_cell_margins(cell, top: int = 80, start: int = 90, bottom: int = 80, end: int = 90) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)

def clear_cell_shading(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is not None:
        shd.set(qn("w:fill"), "auto")
        shd.set(qn("w:val"), "clear")


def set_cell_bottom_border(cell, color: str = "000000", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "0")
    bottom.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color: str = "000000", size: str = "8") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        visible = edge in {"top", "bottom"}
        node.set(qn("w:val"), "single" if visible else "nil")
        node.set(qn("w:sz"), size if visible else "0")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def roman_numeral(value: int) -> str:
    numerals = (
        (10, "X"),
        (9, "IX"),
        (5, "V"),
        (4, "IV"),
        (1, "I"),
    )
    result = []
    for number, symbol in numerals:
        while value >= number:
            result.append(symbol)
            value -= number
    return "".join(result)


def set_fixed_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_twips = int(sum(widths) * 1440)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(int(widths[idx] * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def clear_template_body(doc: Document) -> None:
    body = doc._element.body
    first_sect_pr = deepcopy(doc.sections[0]._sectPr)
    for child in list(body):
        body.remove(child)
    body.append(first_sect_pr)
    set_columns(doc.sections[0], 1)


def set_run_font(run, size: float | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def add_inline_text(paragraph, text: str, size: float = 10.0) -> None:
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            set_run_font(run, size)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size, bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=8.5)
            run.font.name = "Courier New"
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="PARA")
    paragraph.paragraph_format.widow_control = True
    add_inline_text(paragraph, text)


def add_list_item(doc: Document, text: str, marker: str) -> None:
    paragraph = doc.add_paragraph(style="PARA_Indent")
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    marker_run = paragraph.add_run(f"{marker} ")
    set_run_font(marker_run, 10.0, bold=False)
    add_inline_text(paragraph, text)


def add_main_heading(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph(style="H1_List (Space)")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    run.font.color.rgb = IEEE_BLUE


def add_subheading(doc: Document, title: str, letter: str) -> None:
    paragraph = doc.add_paragraph(style="H2_Cont (No Space)")
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(8)
    num_pr = OxmlElement("w:numPr")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "0")
    num_pr.append(num_id)
    paragraph._p.get_or_add_pPr().append(num_pr)
    run = paragraph.add_run(f"{letter}. {title}")
    run.font.color.rgb = GRAY


def add_unnumbered_heading(doc: Document, title: str) -> None:
    paragraph = doc.add_paragraph(style="H1")
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(title)
    run.font.color.rgb = IEEE_BLUE


def start_continuous_section(doc: Document, columns: int) -> None:
    section = doc.add_section(WD_SECTION.CONTINUOUS)
    previous = doc.sections[-2]
    section.page_width = previous.page_width
    section.page_height = previous.page_height
    section.top_margin = previous.top_margin
    section.bottom_margin = previous.bottom_margin
    section.left_margin = previous.left_margin
    section.right_margin = previous.right_margin
    section.header_distance = previous.header_distance
    section.footer_distance = previous.footer_distance
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True
    set_columns(section, columns)


def add_full_width_figure(doc: Document, path: Path, caption: str) -> None:
    start_continuous_section(doc, 1)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(5)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.95))
    cap = doc.add_paragraph(style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.keep_with_next = False
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, 8.0)
    start_continuous_section(doc, 2)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    index = start
    while index < len(lines) and lines[index].strip().startswith("|"):
        line = lines[index].strip()
        if not re.fullmatch(r"\|[\s:\-|]+\|", line):
            rows.append([cell.strip() for cell in line.strip("|").split("|")])
        index += 1
    return rows, index


def table_widths(rows: list[list[str]]) -> list[float]:
    columns = len(rows[0])
    if columns == 2:
        return [2.0, 4.9]
    if columns == 3:
        return [0.7, 1.9, 4.3]
    if columns == 4:
        return [1.7, 1.5, 2.0, 1.7]
    if columns == 5:
        return [3.1, 0.95, 0.95, 0.95, 0.95]
    if columns == 7:
        return [2.1, 1.0, 0.55, 0.9, 0.7, 0.8, 0.8]
    if columns == 11:
        return [1.45] + [0.545] * 10
    return [6.9 / columns] * columns


def add_full_width_table(doc: Document, rows: list[list[str]], number: int) -> None:
    start_continuous_section(doc, 1)
    if number == 6:
        doc.add_page_break()
    leading = doc.add_paragraph()
    leading_run = leading.add_run("\u00A0")
    set_run_font(leading_run, 8.0)
    leading.paragraph_format.line_spacing = 1.0
    leading.paragraph_format.space_before = Pt(0)
    leading.paragraph_format.space_after = Pt(0)
    title = doc.add_paragraph(style="Table Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_text = TABLE_TITLES[number - 1] if number <= len(TABLE_TITLES) else f"SUPPORTING RESULTS {number}"
    title_run = title.add_run(f"TABLE {roman_numeral(number)}\n{title_text}")
    set_run_font(title_run, 8.0, bold=False)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Normal Table"
    set_table_borders(table)
    set_fixed_table_width(table, table_widths(rows))
    set_repeat_table_header(table.rows[0])

    columns = len(rows[0])
    body_size = 6.5 if columns >= 7 else 8.0
    for row_index, values in enumerate(rows):
        for column_index, value in enumerate(values):
            cell = table.cell(row_index, column_index)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            clear_cell_shading(cell)
            if row_index == 0:
                set_cell_bottom_border(cell)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.alignment = (
                WD_ALIGN_PARAGRAPH.LEFT
                if column_index == 0 or len(value) > 24
                else WD_ALIGN_PARAGRAPH.CENTER
            )
            run = paragraph.add_run(value)
            set_run_font(run, body_size, bold=(row_index == 0))
    trailing = doc.add_paragraph()
    trailing_run = trailing.add_run("\u00A0")
    set_run_font(trailing_run, 8.0)
    trailing.paragraph_format.line_spacing = 1.0
    trailing.paragraph_format.space_after = Pt(0)
    start_continuous_section(doc, 2)


def add_algorithm(doc: Document) -> None:
    start_continuous_section(doc, 1)
    title = doc.add_paragraph()
    title.paragraph_format.keep_with_next = True
    title_run = title.add_run("Algorithm 1  Context-Aware Multimodal Re-Ranking in G4")
    set_run_font(title_run, 9.0, bold=True)

    lines = [
        "Input: query q; G3 text candidates T; G3 image candidates I; stage profiles S; context map C",
        "Output: re-ranked Top-k text candidates T' and image candidates I'",
        "1: Encode q using BGE-M3 and compute stage similarities a1 >= a2.",
        "2: If a1 < tau_s or (a1 - a2) < tau_m, return the original G3 Top-k results.",
        "3: Load page ranges, section terms, keywords, and stage weight for the inferred stage.",
        "4: Expand I with valid image candidates from the inferred stage page range.",
        "5: For each text candidate t:",
        "6:     c_t = w_s(0.55 p_t + 0.30 k_t + 0.15 h_t).",
        "7:     f_t = [1 - (rank_G3(t) - 1) / |T|] + beta_t c_t.",
        "8: For each image candidate i:",
        "9:     b_i = 0.25 v_i + 0.15 l_i + 0.50 n_i + 0.05 m_i + 0.05 d_i.",
        "10:    c_i = w_s(0.50 p_i + 0.10 k_i + 0.40 h_i).",
        "11:    Compute rho_i from the baseline rank window R.",
        "12:    f_i = b_i + lambda_c c_i rho_i + lambda_p p_i.",
        "13: Sort candidates by f_t and f_i in descending order and return Top-k.",
        "Parameters: tau_s=0.45, tau_m=0.03, beta_t=0.28, lambda_c=0.50, lambda_p=0.25, R=120.",
        "Symbols: p, k, h denote page, keyword, and section scores; v, l, n, m, d denote image-only, linked-text rank, page proximity, stored SigLIP mapping, and diagram scores.",
        "Leakage control: C contains no query IDs, ground-truth image filenames, or ground-truth chunk IDs.",
    ]

    table = doc.add_table(rows=1, cols=1)
    table.style = "Normal Table"
    set_table_borders(table)
    set_fixed_table_width(table, [6.9])
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=110, start=140, bottom=110, end=140)
    shade_cell(cell, "F7F9FB")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    for index, line in enumerate(lines):
        if index:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        set_run_font(run, 7.6)
        run.font.name = "Courier New"

    note = doc.add_paragraph(style="PARA")
    note.paragraph_format.space_before = Pt(5)
    note.paragraph_format.space_after = Pt(8)
    note_run = note.add_run(
        "The parameter values are linked to the limited grid search described in Section IV-D. "
        "They are fixed across all queries but were selected on the same internal 70-query pilot set used for final evaluation."
    )
    set_run_font(note_run, 8.0, italic=True)

    guide_rows = [
        ("q, T, I", "User query and the G3 text and image candidate sets."),
        ("S, C", "Training-stage profiles and the manual-derived context map."),
        ("a1, a2, tau_s, tau_m", "The two highest stage-profile similarities and their minimum score and margin thresholds."),
        ("w_s, p_x, k_x, h_x", "Stage weight and normalized page-range, keyword, and section-heading match scores. The page score is 1.00 inside the mapped range, 0.70 at one page, 0.45 at two pages, and 0 otherwise."),
        ("v_i, l_i, n_i, m_i, d_i", "Image-only retrieval, linked-text rank, page proximity, stored SigLIP mapping after BBox filtering, and diagram-confidence scores."),
        ("b_x, c_x, f_x", "Baseline, stage-context, and final candidate scores."),
        (
            "beta_t, lambda_c, lambda_p, R, rho_i",
            "Context coefficients; baseline-rank window and rank-dependent attenuation factor.",
        ),
    ]
    guide = doc.add_table(rows=2 + len(guide_rows), cols=2)
    guide.style = "Normal Table"
    set_fixed_table_width(guide, [1.65, 5.25])
    set_table_borders(guide, color="6B7280", size="6")

    title_cell = guide.cell(0, 0).merge(guide.cell(0, 1))
    shade_cell(title_cell, "E9EFF5")
    title_paragraph = title_cell.paragraphs[0]
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_paragraph.paragraph_format.space_after = Pt(0)
    title_run = title_paragraph.add_run("SCORE-TERM GUIDE FOR ALGORITHM 1")
    set_run_font(title_run, 8.5, bold=True)

    for column, label in enumerate(("Term(s)", "Operational meaning")):
        cell = guide.cell(1, column)
        set_cell_margins(cell, top=70, start=90, bottom=70, end=90)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(label)
        set_run_font(run, 7.8, bold=True)
    set_cell_bottom_border(guide.cell(1, 0), color="6B7280", size="6")
    set_cell_bottom_border(guide.cell(1, 1), color="6B7280", size="6")

    for row_index, (symbols, meaning) in enumerate(guide_rows, start=2):
        for column, value in enumerate((symbols, meaning)):
            cell = guide.cell(row_index, column)
            set_cell_margins(cell, top=58, start=90, bottom=58, end=90)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            set_run_font(run, 7.4, italic=(column == 0))

    leakage = doc.add_paragraph(style="PARA")
    leakage.paragraph_format.space_before = Pt(6)
    leakage.paragraph_format.space_after = Pt(2)
    leakage_run = leakage.add_run(
        "All match scores are normalized to [0,1]. If the stage-confidence conditions are not met, G4 returns the original G3 ranking. "
        "For leakage control, context map C contains no query IDs, ground-truth image filenames, or ground-truth chunk identifiers."
    )
    set_run_font(leakage_run, 8.0)
    start_continuous_section(doc, 2)


def extract_front_matter(lines: list[str]) -> tuple[str, str, list[str]]:
    title = "A Context-Aware Multimodal Retrieval-Augmented Generation Framework for Collaborative Robot Training"
    abstract_parts: list[str] = []
    keywords = ""
    mode = ""
    for line in lines:
        stripped = line.strip()
        if stripped == "## Abstract":
            mode = "abstract"
            continue
        if stripped == "## Keywords":
            mode = "keywords"
            continue
        if stripped.startswith("## 1. Introduction"):
            break
        if mode == "abstract" and stripped:
            abstract_parts.append(stripped)
        elif mode == "keywords" and stripped:
            keywords = stripped
    return title, " ".join(abstract_parts), [item.strip() for item in keywords.split(",")]


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def validate_metric_sources(source_text: str) -> None:
    """Fail before rendering when manuscript metrics drift from saved evidence."""
    final_rows = {row["비교군"]: row for row in read_csv_rows(FINAL_METRICS)}
    for group in ("G3", "G4"):
        row = final_rows[group]
        required = (
            row["Image Recall@1"],
            row["Image Recall@5"],
            row["Image Recall@10"],
            row["Image MRR"],
        )
        if not all(value in source_text for value in required):
            raise ValueError(f"{group} 최종 이미지 지표가 영문 원고와 일치하지 않습니다: {required}")

    for row in read_csv_rows(MAPPING_ABLATION):
        table_row = (
            f"| {row['configuration']} | {row['Image Recall@1']} | {row['Image Recall@5']} | "
            f"{row['Image Recall@10']} | {row['Image MRR']} |"
        )
        if table_row not in source_text:
            raise ValueError(f"Table III 수치가 재현 평가와 일치하지 않습니다: {table_row}")

    bootstrap_rows = read_csv_rows(BOOTSTRAP_RESULTS)
    if len(bootstrap_rows) != 4 or any(row["query_count"] != "70" for row in bootstrap_rows):
        raise ValueError("paired bootstrap 근거 파일의 행 수 또는 질의 수가 올바르지 않습니다.")


def validate_reference_citation_order(source_text: str) -> None:
    """Ensure IEEE references are numbered in first-citation order."""
    body_text, references_text = source_text.split("## References", maxsplit=1)
    first_seen: list[int] = []
    for group in re.findall(r"\[([0-9]+(?:\s*,\s*[0-9]+)*)\]", body_text):
        numbers = [int(value.strip()) for value in group.split(",")]
        if 0 in numbers:
            continue
        for number in numbers:
            if number not in first_seen:
                first_seen.append(number)

    expected_citations = list(range(1, len(first_seen) + 1))
    if first_seen != expected_citations:
        raise ValueError(
            "참고문헌 번호가 본문 최초 인용 순서와 일치하지 않습니다: "
            f"{first_seen}"
        )

    reference_numbers = [
        int(value)
        for value in re.findall(r"^\[(\d+)\]\s", references_text, flags=re.MULTILINE)
    ]
    expected_references = list(range(1, len(reference_numbers) + 1))
    if reference_numbers != expected_references:
        raise ValueError(f"References 번호가 연속적이지 않습니다: {reference_numbers}")


def add_front_matter(doc: Document, title: str, abstract: str, keywords: list[str]) -> None:
    dop = doc.add_paragraph(style="DOP")
    dop.add_run("Date of publication xxxx 00, 0000, date of current version xxxx 00, 0000.")
    doi = doc.add_paragraph(style="DOI")
    doi.add_run("Digital Object Identifier 10.1109/ACCESS.2024.Doi Number")

    title_paragraph = doc.add_paragraph(style="Paper Title")
    title_paragraph.add_run(title)

    author = doc.add_paragraph(style="AU")
    author.add_run(AUTHOR_NAME)
    jimin_affiliation = doc.add_paragraph(style="PI_No Space")
    jimin_affiliation.add_run(
        f"Jimin Lee is with {JIMIN_AFFILIATION} "
        f"(e-mail: {AUTHOR_EMAIL}; ORCID: {AUTHOR_ORCID})."
    )
    corresponding_affiliation = doc.add_paragraph(style="PI_No Space")
    corresponding_affiliation.add_run(
        f"Hyun Jung Kim is with {HYUN_JUNG_KIM_AFFILIATION} "
        f"(e-mail: {CORRESPONDING_AUTHOR_EMAIL}; ORCID: {CORRESPONDING_AUTHOR_ORCID})."
    )
    corresponding = doc.add_paragraph(style="PI")
    corresponding.add_run(
        f"Corresponding author: {CORRESPONDING_AUTHOR_NAME} "
        f"(e-mail: {CORRESPONDING_AUTHOR_EMAIL})."
    )

    abstract_paragraph = doc.add_paragraph(style="Abstract")
    label = abstract_paragraph.add_run("ABSTRACT  ")
    label.bold = True
    label.font.color.rgb = IEEE_BLUE
    add_inline_text(abstract_paragraph, abstract)

    index_paragraph = doc.add_paragraph(style="IT")
    label = index_paragraph.add_run("INDEX TERMS  ")
    label.bold = True
    label.font.color.rgb = IEEE_BLUE
    add_inline_text(index_paragraph, ", ".join(keywords))

    start_continuous_section(doc, 2)


def add_references(doc: Document, lines: list[str]) -> None:
    add_unnumbered_heading(doc, "REFERENCES")
    for line in lines:
        paragraph = doc.add_paragraph(style="References")
        paragraph.paragraph_format.keep_together = "[Online]. Available:" in line
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = paragraph.add_run(re.sub(r"^\[\d+\]\s*", "", line))
        set_run_font(run, 8.0)


def add_acknowledgment(doc: Document) -> None:
    add_unnumbered_heading(doc, "ACKNOWLEDGMENT")
    add_body_paragraph(
        doc,
        "The authors thank the Doosan Robotics Marketing Team for confirming that the DART-Platform manual may be included in this research with clear source attribution and without a separate approval procedure. "
        "The manual is cited in [4]. OpenAI Codex [31] was used to assist with code development, data-analysis scripting, figure formatting, and drafting and language revision across Sections I-VIII, the tables, figure captions, and Algorithm 1. "
        "The authors are responsible for all technical decisions, experimental results, citations, and interpretations in the final manuscript."
    )


def add_author_biographies(doc: Document) -> None:
    add_unnumbered_heading(doc, "AUTHOR BIOGRAPHIES")
    add_body_paragraph(doc, AUTHOR_BIOGRAPHY)
    add_body_paragraph(doc, CORRESPONDING_AUTHOR_BIOGRAPHY)


def build() -> Path:
    for required in (
        SOURCE,
        TEMPLATE,
        FIG1,
        FIG2,
        FINAL_METRICS,
        MAPPING_ABLATION,
        BOOTSTRAP_RESULTS,
    ):
        if not required.exists():
            raise FileNotFoundError(required)

    source_text = SOURCE.read_text(encoding="utf-8")
    validate_metric_sources(source_text)
    validate_reference_citation_order(source_text)
    lines = source_text.splitlines()
    title, abstract, keywords = extract_front_matter(lines)

    doc = Document(TEMPLATE)
    clear_template_body(doc)
    ensure_dynamic_page_number(doc)
    add_front_matter(doc, title, abstract, keywords)

    body_start = next(i for i, line in enumerate(lines) if line.startswith("## 1. Introduction"))
    references_start = next(i for i, line in enumerate(lines) if line.strip() == "## References")
    body_lines = lines[body_start:references_start]
    reference_lines = [line.strip() for line in lines[references_start + 1 :] if re.match(r"^\[\d+\]\s", line.strip())]

    current_main = ""
    subsection_index = 0
    current_subsection = ""
    figure1_inserted = False
    figure2_inserted = False
    algorithm_inserted = False
    table_number = 0
    in_code_block = False
    index = 0

    while index < len(body_lines):
        stripped = body_lines[index].strip()
        if not stripped:
            index += 1
            continue

        if stripped.startswith("```"):
            in_code_block = not in_code_block
            index += 1
            continue
        if in_code_block:
            index += 1
            continue

        if stripped.startswith("## "):
            if current_subsection.startswith("3.3") and not figure1_inserted:
                add_full_width_figure(
                    doc,
                    FIG1,
                    "FIGURE 1. Overall architecture of the proposed context-aware multimodal RAG framework.",
                )
                figure1_inserted = True
            current_main = stripped[3:].strip()
            current_subsection = ""
            subsection_index = 0
            fallback_heading = re.sub(r"^\d+\.\s*", "", current_main).upper()
            add_main_heading(doc, MAIN_HEADINGS.get(current_main, fallback_heading))
            index += 1
            continue

        if stripped.startswith("### "):
            if current_subsection.startswith("3.3") and not figure1_inserted:
                add_full_width_figure(
                    doc,
                    FIG1,
                    "FIGURE 1. Overall architecture of the proposed context-aware multimodal RAG framework.",
                )
                figure1_inserted = True
            current_subsection = stripped[4:].strip()
            subsection_index += 1
            add_subheading(doc, re.sub(r"^\d+\.\d+\s*", "", current_subsection), chr(64 + subsection_index))
            index += 1
            continue

        if stripped.startswith("**Algorithm 1"):
            if not figure2_inserted:
                add_full_width_figure(
                    doc,
                    FIG2,
                    "FIGURE 2. Context-aware candidate expansion and re-ranking process used in G4.",
                )
                figure2_inserted = True
            if not algorithm_inserted:
                add_algorithm(doc)
                algorithm_inserted = True
            index += 1
            continue

        if algorithm_inserted and stripped.startswith(
            (
                "In Algorithm 1,",
                "For a candidate x,",
                "If a_1 is below tau_s",
            )
        ):
            # The compact score-term guide emitted by add_algorithm replaces
            # the older three-paragraph notation explanation.
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(body_lines, index)
            table_number += 1
            add_full_width_table(doc, rows, table_number)
            continue

        numbered = re.match(r"^(\d+)\.\s+(.+)", stripped)
        if numbered:
            add_list_item(doc, numbered.group(2), marker=f"{numbered.group(1)})")
            index += 1
            continue
        if stripped.startswith("- "):
            add_list_item(doc, stripped[2:], marker="\u2022")
            index += 1
            continue

        add_body_paragraph(doc, stripped)
        index += 1

    if not figure1_inserted:
        add_full_width_figure(
            doc,
            FIG1,
            "FIGURE 1. Overall architecture of the proposed context-aware multimodal RAG framework.",
        )
    if not figure2_inserted:
        add_full_width_figure(
            doc,
            FIG2,
            "FIGURE 2. Context-aware candidate expansion and re-ranking process used in G4.",
        )
    if not algorithm_inserted:
        add_algorithm(doc)

    add_acknowledgment(doc)
    add_references(doc, reference_lines)
    add_author_biographies(doc)

    doc.core_properties.title = title
    doc.core_properties.author = AUTHOR_NAME
    doc.core_properties.subject = "IEEE Access integrated manuscript"
    doc.core_properties.keywords = ", ".join(keywords)
    update_fields = doc.settings.element.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        doc.settings.element.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    remove_paragraph_pagination_marks(doc)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
